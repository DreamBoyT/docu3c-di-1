import streamlit as st
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
import tempfile
import time
import concurrent.futures
from langchain import LLMChain, PromptTemplate
from langchain.document_loaders import PyPDFLoader
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from PyPDF2 import PdfReader
from io import BytesIO
from docx import Document as DocxDocument
import base64
import re
import nltk
from nltk.corpus import stopwords
from gensim import corpora
from gensim.models.ldamodel import LdaModel
from langchain.text_splitter import CharacterTextSplitter

nltk.download('stopwords')

# Azure OpenAI API details
azure_api_key = 'c09f91126e51468d88f57cb83a63ee36'
azure_endpoint = 'https://chat-gpt-a1.openai.azure.com/'
azure_api_version = '2024-02-01'
azure_chat_endpoint = 'https://danielingitaraj.openai.azure.com/'
openai_api_key = 'a5c4e09a50dd4e13a69e7ef19d07b48c'

# Initialize Azure OpenAI Embeddings
embed_model = AzureOpenAIEmbeddings(
    model="text-embedding-3-large",
    deployment="text-embedding-3-large",
    api_version="2023-12-01-preview",
    azure_endpoint=azure_endpoint,
    openai_api_key=azure_api_key,
)

# Initialize Azure OpenAI LLM
llm = AzureChatOpenAI(
    openai_api_key=openai_api_key,
    api_version=azure_api_version,
    azure_endpoint=azure_chat_endpoint,
    model="gpt-4",
    base_url=None,
    azure_deployment="GPT4",
    temperature=0.5,  # Adjusted temperature for improved summaries
)

# Streamlit user interface
st.title("Document Intelligent Application")
pdf_file = st.file_uploader("Choose a PDF file", type="pdf")

def preprocess_text(text):
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\[.*?\]', '', text)
    text = re.sub(r'\w*\d\w*', '', text)
    text = re.sub(r'http\S+', '', text)
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    return text

def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return [(i + 1, preprocess_text(page.extract_text())) for i, page in enumerate(reader.pages)]

def extract_topics(texts, num_topics=4, num_words=5):
    stop_words = set(stopwords.words('english'))
    texts = [[word for word in text.split() if word not in stop_words] for _, text in texts]
    dictionary = corpora.Dictionary(texts)
    corpus = [dictionary.doc2bow(text) for text in texts]
    lda = LdaModel(corpus, num_topics=num_topics, id2word=dictionary, passes=10)
    topics = lda.print_topics(num_words=num_words)
    return topics

def create_prompt(page_numbers, combined_text, topics):
    combined_text = combined_text.replace("{", "{{").replace("}", "}}")
    topics_text = "\n".join([f"Topic {i+1}: {topic}" for i, topic in enumerate(topics)])
    
    return f"""
    For each of the key topics extracted from the text, please provide the following information in a clear and structured format:

    1. Topic: Provide the topic name in one line. Ensure that the topic is insightful and relevant to the content of the pages.
    
    2. Subtopics: List 3-5 subtopics each that describe the main content of these pages in a line without any extra points added to the subtopic.

    3. Summary: Write a detailed paragraph summary of the topic.

    Here are the topics extracted and the text for reference:

    Topics extracted:
    {topics_text}

    Text:
    {combined_text}
    """

def summarize_pages(llm, page_numbers, combined_text, topics):
    prompt = create_prompt(page_numbers, combined_text, topics)
    prompt_template = PromptTemplate.from_template(prompt)
    chain = LLMChain(llm=llm, prompt=prompt_template)
    response = chain.run({"combined_text": combined_text})
    start_page, end_page = page_numbers[0], page_numbers[-1]
    return f"Pages {start_page}-{end_page}:\n\n{response.strip()}\n"

def group_texts(texts, group_size):
    grouped_texts = []
    for i in range(0, len(texts), group_size):
        group = texts[i:i + group_size]
        page_numbers = [num for num, _ in group]
        combined_text = "\n".join([text for _, text in group])
        grouped_texts.append((page_numbers, combined_text))
    return grouped_texts

def extract_summaries_from_pdf(llm, file, group_size):
    texts = extract_text_from_pdf(file)
    grouped_texts = group_texts(texts, group_size)
    summaries = [None] * len(grouped_texts)

    def process_group(idx, page_numbers, combined_text):
        try:
            topics = extract_topics([(num, combined_text) for num in page_numbers])
            return summarize_pages(llm, page_numbers, combined_text, topics)
        except Exception as e:
            start_page, end_page = page_numbers[0], page_numbers[-1]
            return f"Pages {start_page}-{end_page}:\n\nError summarizing pages {start_page}-{end_page}: {e}\n"

    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = {executor.submit(process_group, idx, page_numbers, combined_text): idx for idx, (page_numbers, combined_text) in enumerate(grouped_texts)}
        for future in concurrent.futures.as_completed(futures):
            idx = futures[future]
            summaries[idx] = future.result()

    return "\n".join(summaries)

def generate_word_file(topics_data):
    doc = DocxDocument()
    doc.add_heading('Document Topics and Subtopics', 0)
    
    for topic_data in topics_data:
        doc.add_heading(topic_data['topic'], level=1)
        doc.add_heading('Subtopics', level=2)
        for i, subtopic in enumerate(topic_data['subtopics']):
            # Remove any leading numbering from the subtopic text
            subtopic = re.sub(r'^\d+\.\s*', '', subtopic)
            doc.add_paragraph(f"{i+1}. {subtopic}", style='List Number')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_topic_data(summary_text):
    topics_data = []
    topic_pattern = re.compile(r'Topic \d+: (.+)')
    subtopics_pattern = re.compile(r'Subtopics:\n(.*?)\n\n', re.DOTALL)
    summary_pattern = re.compile(r'Summary:\n(.*?)\n\n', re.DOTALL)

    topics = topic_pattern.findall(summary_text)
    subtopics = subtopics_pattern.findall(summary_text)
    summaries = summary_pattern.findall(summary_text)

    # Ensure the lists have the same length
    min_length = min(len(topics), len(subtopics), len(summaries))
    topics = topics[:min_length]
    subtopics = subtopics[:min_length]
    summaries = summaries[:min_length]

    for i in range(min_length):
        topics_data.append({
            'topic': topics[i],
            'subtopics': subtopics[i].strip().split('\n'),
            'summary': summaries[i].strip()
        })
    return topics_data

if pdf_file is not None:
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        tmp_file.write(pdf_file.read())
        pdf_path = tmp_file.name
        loader = PyPDFLoader(pdf_path)
        pages = loader.load_and_split()

    summary_option = "Generate 3 Page Summary (default)"
    overall_summary = extract_summaries_from_pdf(llm, pdf_path, group_size=3)

    if overall_summary:
        topics_data = extract_topic_data(overall_summary)
        
        with st.container():
            st.subheader(pdf_file.name)
            for topic_data in topics_data:
                with st.expander(topic_data['topic']):
                    st.write("Subtopics:")
                    for i, subtopic in enumerate(topic_data['subtopics']):
                        # Remove any leading numbering from the subtopic text
                        subtopic = re.sub(r'^\d+\.\s*', '', subtopic)
                        st.write(f"{i+1}. {subtopic}")
                    st.write("Summary:")
                    st.write(topic_data['summary'])

        with st.container():
            col1, col2 = st.columns([3, 1])
            with col1:
                pass
            with col2:
                word_file = generate_word_file(topics_data)
                file_name = pdf_file.name.rsplit('.', 1)[0] + "_topics.docx"
                st.markdown(
                    f"""
                    <style>
                    .download-button {{
                        position: relative;
                        display: inline-block;
                        margin-top: 10px;
                    }}
                    .download-button button {{
                        background-color: #1a1c23;
                        color: white;
                        border: 2px solid #262730;
                        border-radius: 8px;
                        padding: 10px 20px;
                        cursor: pointer;
                        font-size: 14px;
                        transition: background-color 0.3s, border-color 0.3s;
                    }}
                    .download-button button:hover {{
                        border-color: #1a1c23;
                    }}
                    .tooltip {{
                        visibility: hidden;
                        width: 160px;
                        background-color: #555;
                        color: #fff;
                        text-align: center;
                        border-radius: 5px;
                        padding: 5px 0;
                        position: absolute;
                        z-index: 1;
                        bottom: 125%;
                        left: 50%;
                        margin-left: -80px;
                        opacity: 0;
                        transition: opacity 0.3s;
                    }}
                    .download-button:hover .tooltip {{
                        visibility: visible;
                        opacity: 1;
                    }}
                    </style>
                    <div class="download-button">
                        <a href="data:application/octet-stream;base64,{base64.b64encode(word_file.getvalue()).decode()}" download="{file_name}">
                            <button>Download Topics</button>
                            <span class="tooltip" style="background-color: #1a1c23;">Download the topics and subtopics in Word format</span>
                        </a>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
